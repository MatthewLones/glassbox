"""File processor worker - extracts text and generates embeddings."""

import asyncio
import io
import signal
import sys
import tempfile
from pathlib import Path
from typing import Any, Optional

import structlog

sys.path.insert(0, str(__file__).rsplit("/", 2)[0])

from shared.config import get_settings
from shared.db import get_db
from shared.s3 import S3Client
from shared.sqs import SQSConsumer

logger = structlog.get_logger()


async def process_file(file_id: str) -> None:
    """Process a file: extract text and generate embeddings."""
    db = await get_db()

    # Get file info
    file = await db.fetchrow(
        "SELECT * FROM files WHERE id = $1",
        file_id,
    )

    if not file:
        logger.error("File not found", file_id=file_id)
        return

    logger.info(
        "Processing file",
        file_id=file_id,
        filename=file["filename"],
        content_type=file["content_type"],
    )

    try:
        # Update status to processing
        await db.execute(
            "UPDATE files SET processing_status = 'processing' WHERE id = $1",
            file_id,
        )

        # Extract text based on content type
        content_type = file["content_type"] or ""
        extracted_text = ""

        if "pdf" in content_type:
            extracted_text = await extract_pdf(file)
        elif "word" in content_type or "docx" in content_type:
            extracted_text = await extract_docx(file)
        elif "text" in content_type:
            extracted_text = await extract_text(file)
        elif "image" in content_type:
            extracted_text = await extract_image_ocr(file)
        else:
            logger.warning(
                "Unsupported content type",
                content_type=content_type,
                file_id=file_id,
            )

        # Generate embedding
        embedding = None
        if extracted_text:
            embedding = await generate_embedding(extracted_text)

        # Update file record
        if embedding:
            # Format embedding as pgvector string: [0.1, 0.2, ...]
            embedding_str = "[" + ",".join(str(x) for x in embedding) + "]"
            await db.execute(
                """
                UPDATE files
                SET processing_status = 'complete',
                    extracted_text = $1,
                    embedding = $2::vector
                WHERE id = $3
                """,
                extracted_text[:50000],  # Limit text size
                embedding_str,
                file_id,
            )
        else:
            await db.execute(
                """
                UPDATE files
                SET processing_status = 'complete',
                    extracted_text = $1
                WHERE id = $2
                """,
                extracted_text[:50000],
                file_id,
            )

        logger.info(
            "File processed successfully",
            file_id=file_id,
            text_length=len(extracted_text),
            has_embedding=embedding is not None,
        )

    except Exception as e:
        logger.error("File processing failed", file_id=file_id, error=str(e))
        await db.execute(
            """
            UPDATE files
            SET processing_status = 'failed',
                processing_error = $1
            WHERE id = $2
            """,
            str(e),
            file_id,
        )
        raise


async def download_file_from_s3(file: dict) -> bytes:
    """Download file content from S3."""
    s3 = S3Client()
    storage_key = file["storage_key"]

    logger.info("Downloading file from S3", storage_key=storage_key)
    content = await s3.download(storage_key)
    return content


async def extract_pdf(file: dict) -> str:
    """Extract text from PDF using pypdf."""
    from pypdf import PdfReader

    content = await download_file_from_s3(file)

    # Use io.BytesIO to read PDF from bytes
    pdf_file = io.BytesIO(content)
    reader = PdfReader(pdf_file)

    text_parts = []
    for page_num, page in enumerate(reader.pages):
        try:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
        except Exception as e:
            logger.warning(f"Failed to extract page {page_num + 1}", error=str(e))

    full_text = "\n\n".join(text_parts)
    logger.info("Extracted PDF text", pages=len(reader.pages), text_length=len(full_text))
    return full_text


async def extract_docx(file: dict) -> str:
    """Extract text from Word document using python-docx."""
    from docx import Document

    content = await download_file_from_s3(file)

    # Use io.BytesIO to read DOCX from bytes
    docx_file = io.BytesIO(content)
    doc = Document(docx_file)

    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    full_text = "\n\n".join(text_parts)
    logger.info("Extracted DOCX text", paragraphs=len(doc.paragraphs), text_length=len(full_text))
    return full_text


async def extract_text(file: dict) -> str:
    """Extract text from plain text file."""
    content = await download_file_from_s3(file)

    # Try common encodings
    for encoding in ["utf-8", "latin-1", "cp1252"]:
        try:
            text = content.decode(encoding)
            logger.info("Extracted plain text", encoding=encoding, text_length=len(text))
            return text
        except UnicodeDecodeError:
            continue

    # Fallback: decode with errors replaced
    text = content.decode("utf-8", errors="replace")
    logger.warning("Extracted text with replacement characters")
    return text


async def extract_image_ocr(file: dict) -> str:
    """Extract text from image using OCR (placeholder - requires pytesseract)."""
    # Note: Full OCR implementation would require pytesseract and tesseract-ocr installed
    # For now, we'll log that OCR is not implemented and return empty
    logger.warning(
        "OCR not implemented - install pytesseract for image text extraction",
        filename=file["filename"],
    )
    return ""


async def generate_embedding(text: str) -> Optional[list[float]]:
    """Generate embedding using LiteLLM.

    Uses text-embedding-3-small (OpenAI) or voyage-2 (if configured).
    Falls back gracefully if no embedding API is available.
    """
    from litellm import aembedding

    settings = get_settings()

    # Truncate text if too long (embedding models have token limits)
    max_chars = 8000
    if len(text) > max_chars:
        text = text[:max_chars]
        logger.info("Truncated text for embedding", original_length=len(text), max_chars=max_chars)

    # Determine which embedding model to use
    if settings.openai_api_key:
        model = "text-embedding-3-small"
        api_key = settings.openai_api_key
    elif settings.anthropic_api_key:
        # Anthropic doesn't have embeddings, but Voyage AI does (common pairing)
        # For now, skip if no OpenAI key
        logger.warning("No OpenAI API key configured for embeddings - skipping embedding generation")
        return None
    else:
        logger.warning("No embedding API key configured - skipping embedding generation")
        return None

    try:
        response = await aembedding(
            model=model,
            input=text,
            api_key=api_key,
        )
        embedding = response.data[0]["embedding"]
        logger.info("Generated embedding", model=model, dimensions=len(embedding))
        return embedding
    except Exception as e:
        logger.error("Embedding generation failed", error=str(e), model=model)
        return None


async def handle_file_job(message: dict[str, Any]) -> None:
    """Handle a file processing job."""
    file_id = message.get("file_id")
    action = message.get("action", "process")

    if not file_id:
        logger.error("Invalid message: missing file_id")
        return

    if action == "process":
        await process_file(file_id)
    else:
        logger.warning("Unknown action", action=action)


async def main() -> None:
    """Main entry point for the file processor worker."""
    settings = get_settings()

    structlog.configure(
        processors=[
            structlog.processors.TimeStamper(fmt="iso"),
            structlog.processors.JSONRenderer(),
        ]
    )

    logger.info("Starting file processor worker", queue_url=settings.sqs_file_queue_url)

    consumer = SQSConsumer(
        queue_url=settings.sqs_file_queue_url,
        handler=handle_file_job,
        visibility_timeout=300,  # 5 minutes for file processing
    )

    # Handle graceful shutdown
    loop = asyncio.get_event_loop()

    def shutdown():
        logger.info("Received shutdown signal")
        consumer.stop()

    for sig in (signal.SIGTERM, signal.SIGINT):
        loop.add_signal_handler(sig, shutdown)

    try:
        await consumer.start()
    except asyncio.CancelledError:
        logger.info("Worker cancelled")
    finally:
        from shared.db import db
        await db.disconnect()
        logger.info("Worker stopped")


if __name__ == "__main__":
    asyncio.run(main())
